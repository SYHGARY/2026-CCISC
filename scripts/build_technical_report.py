from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_DIR = ROOT / "submission"
SCRATCH_DIR = (
    ROOT
    / "outputs"
    / "019eb232-9213-73b3-b220-ef15fca2a827"
    / "documents"
    / "logic-guard-report"
)
OUTPUT_PATH = SUBMISSION_DIR / "LLM_Logic_Consistency_Guard_Technical_Report.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17212B"
MUTED = "5E6A75"
GREEN = "2F7D6D"
PALE_GREEN = "EAF4F1"
PALE_BLUE = "EAF1F8"
PALE_GRAY = "F4F6F9"
WHITE = "FFFFFF"
RED = "A13D3D"
ORANGE = "C46A20"

FONT_LATIN = "Calibri"
FONT_CJK = "SimSun"
FONT_CJK_HEADING = "SimHei"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_inches: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total_dxa = sum(round(width * 1440) for width in widths_inches)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for index, width in enumerate(widths_inches):
            cell = row.cells[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    latin: str = FONT_LATIN,
    cjk: str = FONT_CJK,
) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cjk)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph_runs(paragraph, *, cjk=FONT_CJK) -> None:
    for run in paragraph.runs:
        set_run_font(run, cjk=cjk)


def add_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.28

    for style_name, size, color, before, after in (
        ("Title", 28, INK, 0, 8),
        ("Heading 1", 16, BLUE, 0, 9),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11.5, DARK_BLUE, 7, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK_HEADING)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.05

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(10.3)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("LLM LOGIC CONSISTENCY GUARD")
    set_run_font(run, size=8.5, bold=True, color=MUTED)
    run = hp.add_run("  |  技术报告")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    run = fp.add_run("比赛提交材料  ·  ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(fp, "PAGE")


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix) :])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, size=10.3)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run, size=10.3)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=8.5, color=MUTED)
    return p


def add_callout(doc: Document, title: str, text: str, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    title_run = p.add_run(title)
    set_run_font(title_run, size=10.5, bold=True, color=accent, cjk=FONT_CJK_HEADING)
    detail = cell.add_paragraph()
    detail.paragraph_format.space_after = Pt(1)
    detail.paragraph_format.line_spacing = 1.2
    run = detail.add_run(text)
    set_run_font(run, size=9.8, color=INK)
    return table


def add_metric_strip(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    set_table_widths(table, [1.625, 1.625, 1.625, 1.625])
    metrics = [
        ("26", "单元测试通过"),
        ("28", "对抗基准轨迹"),
        ("4", "一致性检测类型"),
        ("0", "API Key 依赖"),
    ]
    for cell, (value, label) in zip(table.rows[0].cells, metrics):
        set_cell_shading(cell, PALE_GREEN)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(value)
        set_run_font(run, size=18, bold=True, color=GREEN)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=8.5, color=MUTED)


def add_data_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    *,
    font_size=9.0,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    set_repeat_table_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, PALE_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=font_size, bold=True, color=DARK_BLUE, cjk=FONT_CJK_HEADING)
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            run = p.add_run(value)
            set_run_font(run, size=font_size, color=INK)


def load_font(size: int, bold=False):
    path = Path("C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc")
    return ImageFont.truetype(str(path), size=size)


def rounded_box(draw, xy, fill, outline=None, radius=18, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def build_architecture_image(path: Path) -> None:
    image = Image.new("RGB", (1500, 600), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42, bold=True)
    label_font = load_font(29, bold=True)
    small_font = load_font(22)
    draw.text((56, 35), "可观测轨迹上的检测—干预—复检闭环", font=title_font, fill="#17212B")

    boxes = [
        (55, 155, 285, 360, "轨迹归一化", "目标、计划、工具调用\n工具结果、最终回答"),
        (350, 155, 580, 360, "语义结构化", "显式约束提取\n实体化声明提取"),
        (645, 155, 875, 360, "四类检测器", "目标偏离、计划冲突\n自相矛盾、结果失配"),
        (940, 155, 1170, 360, "分级干预", "审计、自检、修订\n人工确认、阻断"),
        (1235, 155, 1465, 360, "修复后复检", "最多 3 轮\n仅接受证据支持的修复"),
    ]
    fills = ["#EAF1F8", "#F4F6F9", "#FFF3E8", "#FCECEC", "#EAF4F1"]
    accents = ["#2E74B5", "#5E6A75", "#C46A20", "#A13D3D", "#2F7D6D"]
    for index, (x1, y1, x2, y2, title, detail) in enumerate(boxes):
        rounded_box(draw, (x1, y1, x2, y2), fills[index], accents[index], 18, 3)
        draw.text((x1 + 24, y1 + 33), title, font=label_font, fill=accents[index])
        draw.multiline_text(
            (x1 + 24, y1 + 95),
            detail,
            font=small_font,
            fill="#33404C",
            spacing=13,
        )
        if index < len(boxes) - 1:
            x = x2 + 13
            y = 257
            draw.line((x, y, x + 39, y), fill="#81909E", width=5)
            draw.polygon([(x + 39, y), (x + 25, y - 10), (x + 25, y + 10)], fill="#81909E")
    draw.line((1350, 390, 1350, 485, 760, 485, 760, 405), fill="#2F7D6D", width=5)
    draw.polygon([(760, 405), (749, 421), (771, 421)], fill="#2F7D6D")
    draw.text((908, 500), "发现新问题则继续；无剩余冲突则输出审计报告", font=small_font, fill="#2F7D6D")
    image.save(path)


def build_evaluation_image(path: Path) -> None:
    image = Image.new("RGB", (1500, 640), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42, bold=True)
    label_font = load_font(26, bold=True)
    small_font = load_font(22)
    value_font = load_font(26, bold=True)
    draw.text((55, 35), "仓库内基准的可重复验证结果", font=title_font, fill="#17212B")
    labels = ["Precision", "Recall", "Micro-F1", "Macro-F1", "Trace Accuracy"]
    colors = ["#2E74B5", "#2F7D6D", "#C46A20", "#A13D3D", "#586B7E"]
    y = 135
    for label, color in zip(labels, colors):
        draw.text((65, y), label, font=label_font, fill="#33404C")
        rounded_box(draw, (345, y - 5, 1350, y + 38), "#E7EBEF", radius=12, width=0)
        rounded_box(draw, (345, y - 5, 1350, y + 38), color, radius=12, width=0)
        draw.text((1375, y - 2), "1.00", font=value_font, fill=color)
        y += 78
    draw.text(
        (65, 555),
        "28 条手工构造轨迹 · 17 个正类标签 · 95 个真负类标签 · FPR 0.00",
        font=small_font,
        fill="#5E6A75",
    )
    image.save(path)


def add_picture(doc: Document, path: Path, width: float, alt_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)


def build_report() -> None:
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
    SCRATCH_DIR.mkdir(parents=True, exist_ok=True)
    architecture_image = SCRATCH_DIR / "architecture.png"
    evaluation_image = SCRATCH_DIR / "evaluation.png"
    build_architecture_image(architecture_image)
    build_evaluation_image(evaluation_image)

    verification = json.loads((ROOT / "outputs" / "verification_summary.json").read_text(encoding="utf-8"))
    adv_metrics = verification["adversarial_metrics"]

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "LLM Logic Consistency Guard 技术报告"
    doc.core_properties.subject = "比赛项目技术说明与验证报告"
    doc.core_properties.author = "[团队名称待填写]"
    doc.core_properties.keywords = "LLM, Agent, Logic Consistency, Runtime Guard, Evaluation"

    # Page 1: editorial cover.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(72)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("比赛项目技术报告")
    set_run_font(run, size=11, bold=True, color=GREEN, cjk=FONT_CJK_HEADING)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("LLM Logic Consistency Guard")
    set_run_font(run, size=28, bold=True, color=INK, cjk=FONT_CJK_HEADING)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    run = subtitle.add_run("面向可观测智能体轨迹的逻辑一致性检测、分级干预与修复复检")
    set_run_font(run, size=14, bold=True, color=DARK_BLUE, cjk=FONT_CJK_HEADING)

    add_callout(
        doc,
        "核心命题",
        "传统安全检查通常判断“单个动作是否危险”；本项目进一步判断智能体的整条执行轨迹，"
        "是否持续遵守用户目标、可见计划、既有声明与工具事实。",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_metric_strip(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    metadata = doc.add_table(rows=3, cols=2)
    set_table_widths(metadata, [3.25, 3.25])
    metadata.style = "Table Grid"
    metadata_rows = [
        ("竞赛 / 赛道", "[待按官方要求填写]"),
        ("团队 / 成员", "[待填写]"),
        ("报告版本", "通用提交版 · 2026-06-12"),
    ]
    for row, (label, value) in zip(metadata.rows, metadata_rows):
        left, right = row.cells
        set_cell_shading(left, PALE_GRAY)
        left_p = left.paragraphs[0]
        left_p.paragraph_format.space_after = Pt(0)
        run = left_p.add_run(label)
        set_run_font(run, size=9.5, bold=True, color=DARK_BLUE, cjk=FONT_CJK_HEADING)
        right_p = right.paragraphs[0]
        right_p.paragraph_format.space_after = Pt(0)
        run = right_p.add_run(value)
        set_run_font(run, size=9.5, color=INK)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(16)
    run = note.add_run("开源式本地原型 · 无 API Key · 标准库运行时 · 可复现实验")
    set_run_font(run, size=9, color=MUTED)

    # Page 2: problem and value.
    add_page_break(doc)
    doc.add_heading("1. 项目摘要与问题定义", level=1)
    add_body(
        doc,
        "LLM Logic Consistency Guard 是一个运行时一致性守卫原型。它接收用户目标、可见计划或推理摘要、"
        "工具调用、工具结果和最终回答组成的可审计轨迹，在执行过程中识别前后矛盾，并依据置信度与动作"
        "可逆性选择审计、自检、修订、人工确认或阻断。",
    )
    add_callout(
        doc,
        "设计边界",
        "系统不读取也不依赖隐藏思维链。所有判断均来自可记录、可复核、可回放的外部轨迹，"
        "因此适合用作 Agent 中间件、评测器或事后审计组件。",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    doc.add_heading("为什么需要轨迹级一致性", level=2)
    add_bullet(doc, "目标偏离：用户要求只读，智能体随后执行写入、删除或外发。")
    add_bullet(doc, "计划—动作冲突：可见计划承诺“不删除”，后续动作却调用 delete_file。")
    add_bullet(doc, "自相矛盾：同一实体在没有有效状态变化时，被先后声明为存在与不存在。")
    add_bullet(doc, "工具—回答失配：工具返回失败或阻断，最终回答却宣称任务成功。")
    doc.add_heading("项目目标", level=2)
    add_number(doc, "把异构日志归一化为统一、可审计的 LogicTrace 数据契约。")
    add_number(doc, "提供可解释的四类检测器，并给出证据、严重度、置信度和干预建议。")
    add_number(doc, "对证据充分的问题实施确定性修复，并在修复后重新运行完整检测。")
    add_number(doc, "以本地交互演示、命令行脚本和自动验证脚本形成可提交、可复现的闭环。")
    doc.add_heading("相对单动作安全检查的增量价值", level=2)
    add_data_table(
        doc,
        ["比较维度", "单动作检查", "本项目"],
        [
            ["观察范围", "当前动作与参数", "用户目标到最终回答的完整轨迹"],
            ["核心问题", "动作是否危险", "轨迹是否前后一致且有事实支撑"],
            ["输出", "允许 / 拒绝", "证据、类型、置信度、分级干预与复检"],
            ["误判控制", "依赖静态规则", "实体区分、状态变化失效与工具事实优先"],
        ],
        [1.35, 2.15, 3.0],
        font_size=8.8,
    )

    # Page 3: architecture.
    add_page_break(doc)
    doc.add_heading("2. 技术方案与运行流程", level=1)
    add_picture(doc, architecture_image, 6.45, "系统从轨迹归一化到修复后复检的闭环架构图")
    add_caption(doc, "图 1  系统运行时闭环：所有判断基于可观测轨迹")
    doc.add_heading("统一轨迹契约", level=2)
    add_body(
        doc,
        "LogicTrace 以 trace_id 和 user_goal 为入口，每个步骤记录 step_id、角色、可见内容、动作名称、"
        "动作参数、工具结果及来源元数据。security 目录产生的 before/after 事件会按动作配对合并；"
        "原安全标签仅作为来源信息保存，不被复用为逻辑一致性真值。",
    )
    doc.add_heading("关键模块", level=2)
    add_data_table(
        doc,
        ["模块", "职责", "关键输出"],
        [
            ["converters.py", "转换并合并异构运行日志", "LogicTrace"],
            ["claims.py", "抽取主体、谓词、极性、来源与证据", "Claim"],
            ["detector.py", "运行四类轨迹一致性检测", "Violation"],
            ["corrector.py", "依据置信度与可逆性选择干预", "Correction"],
            ["repair.py", "应用确定性修复并最多复检 3 轮", "RepairResult"],
            ["evaluation.py", "计算微观、宏观、分类与轨迹指标", "EvaluationReport"],
        ],
        [1.45, 3.15, 1.9],
        font_size=8.6,
    )
    doc.add_heading("时间与实体语义", level=2)
    add_bullet(doc, "声明以 subject + predicate + polarity 表示，避免把不同文件或不同对象混为一谈。")
    add_bullet(doc, "成功的状态变更动作会使旧状态声明失效，避免把合法变化误报为前后矛盾。")
    add_bullet(doc, "工具观察与最终回答冲突时，优先归类为 tool_result_final_mismatch，避免重复报错。")

    # Page 4: detectors and interventions.
    add_page_break(doc)
    doc.add_heading("3. 检测器与分级干预", level=1)
    add_data_table(
        doc,
        ["检测类型", "判断依据", "典型证据", "默认处置"],
        [
            ["goal_deviation", "动作违反用户显式目标约束", "只读目标 + write/delete/send", "人工确认或阻断"],
            ["plan_action_contradiction", "动作违反可见计划或承诺", "no_delete + delete_file", "人工确认或阻断"],
            ["self_contradiction", "同一实体声明极性反转且无状态变化", "文件存在 → 文件不存在", "自检"],
            ["tool_result_final_mismatch", "失败工具结果与成功型最终回答冲突", "File not found + 已成功读取", "修订回答"],
        ],
        [1.45, 2.15, 1.75, 1.15],
        font_size=8.45,
    )
    doc.add_heading("干预阶梯", level=2)
    add_data_table(
        doc,
        ["级别", "适用条件", "系统行为"],
        [
            ["audit", "低风险或低置信度提示", "记录证据，不改变执行。"],
            ["self_check", "可逆、语义性软冲突", "要求智能体比较证据并自检。"],
            ["revise", "工具事实足以支持确定性纠正", "修订最终回答或可见声明。"],
            ["human_confirm", "状态变更存在冲突但需用户裁决", "暂停执行，展示动作与冲突证据。"],
            ["block", "高置信度且不可逆的明显违反", "执行前阻断，并要求重建安全计划。"],
        ],
        [1.15, 2.35, 3.0],
        font_size=8.8,
    )
    doc.add_heading("修复后复检", level=2)
    add_body(
        doc,
        "修复器不会把“已采取纠正措施”等同于“问题已解决”。每轮干预后都重新执行完整检测；"
        "只有剩余违规为空时才标记 resolved=true。当前循环上限为 3 轮，用于避免无界修复。",
    )
    add_callout(
        doc,
        "确定性修复原则",
        "仅修复能被可观测证据直接支持的问题，例如把“读取成功”改为“文件不存在，任务未完成”。"
        "无法可靠判断真假的软语义冲突保留给自检或人工确认，不进行猜测式改写。",
        fill="FFF3E8",
        accent=ORANGE,
    )
    doc.add_heading("一个完整修复示例", level=2)
    add_data_table(
        doc,
        ["阶段", "轨迹内容"],
        [
            ["冲突", "计划承诺只读，但调用 delete_file(report.txt)。"],
            ["干预", "阻断删除；保留原始动作参数到 metadata。"],
            ["同步修订", "最终回答改为“动作已阻断，任务未成功完成”。"],
            ["复检", "重新运行四类检测器，剩余违规为 0。"],
        ],
        [1.15, 5.35],
        font_size=8.8,
    )

    # Page 5: evaluation.
    add_page_break(doc)
    doc.add_heading("4. 实验设计与验证结果", level=1)
    add_picture(doc, evaluation_image, 6.45, "手工对抗基准中五项指标均为一点零")
    add_caption(doc, "图 2  28 条手工对抗基准的仓库内评估结果")
    add_data_table(
        doc,
        ["评估对象", "规模", "Precision", "Recall", "F1 / Macro-F1", "FPR"],
        [
            ["合成回归集", "8 traces", "1.00", "1.00", "1.00 / 1.00", "0.00"],
            ["手工对抗集", "28 traces", "1.00", "1.00", "1.00 / 1.00", "0.00"],
            ["核心修复样例", "2 inconsistent", "—", "—", "2 / 2 已解决", "—"],
            ["security 转换日志", "8 events / 4 actions", "—", "—", "检测并修复", "—"],
        ],
        [1.35, 1.25, 0.85, 0.75, 1.45, 0.85],
        font_size=8.25,
    )
    source = doc.add_paragraph()
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)
    run = source.add_run(
        f"验证快照：TP={adv_metrics['true_positives']}，FP={adv_metrics['false_positives']}，"
        f"FN={adv_metrics['false_negatives']}，TN={adv_metrics['true_negatives']}；"
        "26 个 unittest 全部通过。"
    )
    set_run_font(run, size=8.5, color=MUTED)
    doc.add_heading("分类覆盖", level=2)
    add_data_table(
        doc,
        ["类别", "支持数", "F1"],
        [
            ["goal_deviation", "5", "1.00"],
            ["plan_action_contradiction", "3", "1.00"],
            ["self_contradiction", "5", "1.00"],
            ["tool_result_final_mismatch", "4", "1.00"],
        ],
        [3.7, 1.4, 1.4],
        font_size=8.8,
    )
    add_callout(
        doc,
        "结果解释边界",
        "两套数据均为仓库内人工构造基准，适合证明功能闭环、回归稳定性和可复现性；"
        "它们不是独立采集或公开标注数据，因此 1.00 指标不代表真实开放环境中的泛化能力。",
        fill="FCECEC",
        accent=RED,
    )

    # Page 6: demo, reproducibility, limitations.
    add_page_break(doc)
    doc.add_heading("5. 演示、复现与后续计划", level=1)
    screenshot = ROOT / "demo_screenshot.png"
    add_picture(doc, screenshot, 5.75, "本地交互式演示界面，展示修复后的轨迹与复检结果")
    add_caption(doc, "图 3  本地演示：冲突动作被阻断，最终回答同步修订，复检后无剩余违规")
    doc.add_heading("一分钟复现", level=2)
    add_number(doc, "运行 python scripts\\verify_project.py，执行测试、两套评估与修复验证。")
    add_number(doc, "运行 python scripts\\demo_server.py，启动本地演示服务。")
    add_number(doc, "浏览器打开 http://127.0.0.1:8765，选择样例并执行分析或修复复检。")
    add_body(doc, "运行条件：Python 3.11+；无需 API Key；核心运行时仅依赖 Python 标准库。")
    doc.add_heading("当前局限", level=2)
    add_bullet(doc, "规则语言覆盖有限，对开放域软语义矛盾仍需 NLI 模型或 LLM Judge 辅助。")
    add_bullet(doc, "尚未在独立公开数据集上进行盲测，也未完成真实生产流量的延迟评估。")
    add_bullet(doc, "与 live agent framework 的集成目前通过日志转换展示，尚未提供通用中间件适配层。")
    doc.add_heading("下一阶段里程碑", level=2)
    add_data_table(
        doc,
        ["优先级", "里程碑", "验收标准"],
        [
            ["P0", "引入独立标注的轨迹一致性测试集", "报告盲测 F1、FPR 与失败案例。"],
            ["P1", "增加模型辅助声明抽取器", "保持确定性基线，可切换并比较增益。"],
            ["P1", "接入真实 Agent 中间件", "测量检测延迟、阻断率和修复成功率。"],
            ["P2", "完善多语言与开放谓词覆盖", "扩展中文、英文及领域动作词表。"],
        ],
        [0.75, 2.7, 3.05],
        font_size=8.5,
    )
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(8)
    closing.paragraph_format.space_after = Pt(0)
    run = closing.add_run("项目状态：核心功能、测试、评估、演示与通用提交材料已完成；官方字段待赛题要求确定后填写。")
    set_run_font(run, size=9.2, bold=True, color=GREEN, cjk=FONT_CJK_HEADING)

    for paragraph in doc.paragraphs:
        style_paragraph_runs(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    style_paragraph_runs(paragraph)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_report()
